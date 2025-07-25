import os
import json
import requests
from docx import Document
from PIL import Image
import pytesseract

# Configurations
OPENAI_API_KEY = "YOUR_AZURE_OPENAI_API_KEY"
OPENAI_ENDPOINT = "YOUR_AZURE_OPENAI_ENDPOINT"  # e.g., https://<resource-name>.openai.azure.com/openai/deployments/<deployment-id>/chat/completions?api-version=2024-02-15-preview
HEADERS = {
    "api-key": OPENAI_API_KEY,
    "Content-Type": "application/json"
}

def extract_text_from_image(image_path):
    """Extract text from image using pytesseract."""
    image = Image.open(image_path)
    text = pytesseract.image_to_string(image, lang='vie')
    return text

def call_openai_azure_api(prompt):
    """Call Azure OpenAI API to extract required information."""
    data = {
        "messages": [
            {"role": "system", "content": "Bạn là một luật sư. Hãy trích xuất thông tin bị đơn, nguyên đơn và vụ việc từ văn bản sau."},
            {"role": "user", "content": prompt}
        ],
        "max_tokens": 500,
        "temperature": 0.2
    }
    response = requests.post(OPENAI_ENDPOINT, headers=HEADERS, json=data)
    response.raise_for_status()
    return response.json()["choices"][0]["message"]["content"]

def parse_openai_response(response_text):
    """Parse OpenAI response to JSON."""
    # Expecting OpenAI to return JSON, else parse manually
    try:
        return json.loads(response_text)
    except Exception:
        # Fallback: simple parsing (customize as needed)
        lines = response_text.split('\n')
        result = {}
        for line in lines:
            if ':' in line:
                key, value = line.split(':', 1)
                result[key.strip()] = value.strip()
        return result

def save_metadata(metadata, json_path):
    with open(json_path, 'w', encoding='utf-8') as f:
        json.dump(metadata, f, ensure_ascii=False, indent=2)

def load_metadata(json_path):
    with open(json_path, 'r', encoding='utf-8') as f:
        return json.load(f)

def export_to_doc(metadata, doc_path):
    doc = Document()
    doc.add_heading('Bộ Hồ Sơ Luật', 0)
    for key, value in metadata.items():
        doc.add_heading(key, level=1)
        doc.add_paragraph(str(value))
    doc.save(doc_path)

def main():
    # Step 1: Extract info from image and save to JSON
    image_path = "../assets/tesst.jpg"
    json_path = "../assets/tesst.json"
    doc_path = "tesst.docx"

    print("Đang trích xuất văn bản từ ảnh...")
    text = extract_text_from_image(image_path)
    print("Gọi API OpenAI Azure để bóc dữ liệu...")
    openai_response = call_openai_azure_api(text)
    metadata = parse_openai_response(openai_response)
    save_metadata(metadata, json_path)
    print(f"Đã lưu metadata vào {json_path}")

    # Step 2: Read JSON and export to DOCX
    print("Đang xuất bộ hồ sơ luật ra file Word...")
    metadata = load_metadata(json_path)
    export_to_doc(metadata, doc_path)
    print(f"Đã xuất bộ hồ sơ luật ra {doc_path}")

if __name__ == "__main__":
    main()